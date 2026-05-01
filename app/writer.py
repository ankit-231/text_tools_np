from docx import Document
from docx.shared import Pt
from docx.document import Document as DocumentObject

FONT_SIZE = 15
# FONT_NAME = "Times New Roman"
FONT_NAME = "Preeti"


def create_doc():
    return Document()


def add_page(doc: DocumentObject, content):
    for line in content.split("\n"):
        para = doc.add_paragraph()
        run = para.add_run(line)

        font = run.font
        font.name = FONT_NAME
        font.size = Pt(FONT_SIZE)

    doc.add_page_break()


def save_doc(doc: DocumentObject, path):
    doc.save(path)


# if __name__ == "__main__":
#     contents = [
#         "This is for page 1, line 1\npage 1 line 2",
#         "This is for page 2",
#         "This is for page 3, line 1\npage 3 line 2\npage 3 line 3",
#     ]
#     doc = create_doc()
#     for content in contents:
#         add_page(doc, content)

#     output_path = "files/output.docx"

#     save_doc(doc, output_path)
