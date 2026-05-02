from docx import Document
from docx.shared import Pt
from docx.document import Document as DocumentObject

NEPALI_FONT = "Preeti"
NEPALI_SIZE = Pt(15)

ENGLISH_FONT = "Times New Roman"
ENGLISH_SIZE = Pt(12)


def create_doc():
    return Document()


def _apply_font(run, is_nepali: bool):
    run.font.name = NEPALI_FONT if is_nepali else ENGLISH_FONT
    run.font.size = NEPALI_SIZE if is_nepali else ENGLISH_SIZE


def add_page(
    doc: DocumentObject, converted_lines: list[tuple[str, list[tuple[str, bool]]]]
):
    """
    converted_lines: list of (converted_text, segments) per line.
        converted_text  - the full converted line (used for reference, not directly)
        segments        - list of (chunk, is_nepali) where chunk is already
                          Preeti-encoded if is_nepali, else original English text
    """
    for _converted_text, segments in converted_lines:
        para = doc.add_paragraph()
        for chunk, is_nepali in segments:
            run = para.add_run(chunk)
            _apply_font(run, is_nepali)

    doc.add_page_break()


def save_doc(doc: DocumentObject, path):
    doc.save(path)


# if __name__ == "__main__":
#     contents = [
#         "This is for page 1, line 1\npage 1 line 2",
#         "This is for page 2",
#         "This is for page 3, line 1\npage 3 line 2\npage 3 line 3",
#     ]
#     doc = create_doc()
#     for content in contents:
#         add_page(doc, content)

#     output_path = "files/output.docx"

#     save_doc(doc, output_path)
